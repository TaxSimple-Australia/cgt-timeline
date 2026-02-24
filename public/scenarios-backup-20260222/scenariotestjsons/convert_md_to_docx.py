"""
Markdown to Word Document Converter - Enhanced Version
Converts ALL_CGT_SCENARIOS.md to a professionally formatted Word document
Handles all markdown elements: tables, headers, lists, links, bold, calculations
"""

import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_shading(cell, color):
    """Set cell background color"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_borders(cell, border_color="000000", border_size="4"):
    """Set cell borders"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), border_size)
        border.set(qn('w:color'), border_color)
        tcBorders.append(border)
    tcPr.append(tcBorders)


def add_horizontal_line(doc):
    """Add a horizontal line paragraph"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run('─' * 100)
    run.font.color.rgb = RGBColor(180, 180, 180)
    run.font.size = Pt(6)


def create_styled_document():
    """Create document with custom styles"""
    doc = Document()

    # Set narrow margins
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Heading 1 - Main title
    h1_style = doc.styles['Heading 1']
    h1_style.font.name = 'Calibri'
    h1_style.font.size = Pt(26)
    h1_style.font.bold = True
    h1_style.font.color.rgb = RGBColor(0, 51, 102)
    h1_style.paragraph_format.space_before = Pt(18)
    h1_style.paragraph_format.space_after = Pt(12)

    # Heading 2 - Scenario titles
    h2_style = doc.styles['Heading 2']
    h2_style.font.name = 'Calibri'
    h2_style.font.size = Pt(18)
    h2_style.font.bold = True
    h2_style.font.color.rgb = RGBColor(0, 82, 136)
    h2_style.paragraph_format.space_before = Pt(18)
    h2_style.paragraph_format.space_after = Pt(8)

    # Heading 3 - Subsections
    h3_style = doc.styles['Heading 3']
    h3_style.font.name = 'Calibri'
    h3_style.font.size = Pt(14)
    h3_style.font.bold = True
    h3_style.font.color.rgb = RGBColor(51, 51, 51)
    h3_style.paragraph_format.space_before = Pt(12)
    h3_style.paragraph_format.space_after = Pt(6)

    return doc


def parse_table(lines, start_idx):
    """Parse markdown table starting at given index"""
    table_lines = []
    idx = start_idx

    while idx < len(lines) and lines[idx].strip().startswith('|'):
        table_lines.append(lines[idx].strip())
        idx += 1

    if len(table_lines) < 2:
        return None, start_idx

    # Parse header
    header = [cell.strip() for cell in table_lines[0].split('|')[1:-1]]

    # Skip separator line
    data_start = 2 if len(table_lines) > 1 and '---' in table_lines[1] else 1

    # Parse data rows
    rows = []
    for line in table_lines[data_start:]:
        if '---' not in line:
            row = [cell.strip() for cell in line.split('|')[1:-1]]
            if row:
                rows.append(row)

    return {'header': header, 'rows': rows}, idx


def add_table_to_doc(doc, table_data):
    """Add a formatted table to the document"""
    if not table_data or not table_data['header']:
        return

    num_cols = len(table_data['header'])
    num_rows = len(table_data['rows']) + 1  # +1 for header

    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = 'Table Grid'

    # Auto-fit table
    table.autofit = True

    # Add header row
    header_row = table.rows[0]
    for idx, cell_text in enumerate(table_data['header']):
        cell = header_row.cells[idx]
        cell.text = cell_text
        set_cell_shading(cell, "003366")
        set_cell_borders(cell, "003366")
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.size = Pt(10)
                run.font.name = 'Calibri'

    # Add data rows
    for row_idx, row_data in enumerate(table_data['rows']):
        row = table.rows[row_idx + 1]
        for col_idx, cell_text in enumerate(row_data):
            if col_idx < len(row.cells):
                cell = row.cells[col_idx]
                cell.text = cell_text
                # Alternate row colors
                if row_idx % 2 == 0:
                    set_cell_shading(cell, "F0F4F8")
                else:
                    set_cell_shading(cell, "FFFFFF")
                set_cell_borders(cell, "CCCCCC")
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    for run in paragraph.runs:
                        run.font.size = Pt(10)
                        run.font.name = 'Calibri'

    # Add spacing after table
    doc.add_paragraph()


def process_inline_formatting(paragraph, text):
    """Process inline bold text and add to paragraph"""
    # Remove markdown links and keep text only
    text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)

    # Handle bold text
    parts = re.split(r'(\*\*[^*]+\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            run = paragraph.add_run(part)
        run.font.name = 'Calibri'
        run.font.size = Pt(11)


def is_result_line(text):
    """Check if line is a RESULT or final answer line"""
    text_upper = text.upper()
    return (
        'RESULT:' in text_upper or
        'NET CAPITAL GAIN' in text_upper or
        'NO CGT PAYABLE' in text_upper or
        'FULLY EXEMPT' in text_upper or
        'TOTAL NET CAPITAL GAIN' in text_upper
    )


def is_step_line(text):
    """Check if line is a Step header"""
    return bool(re.match(r'^Step \d+:', text.strip()))


def is_calculation_header(text):
    """Check if line is a calculation section header"""
    text_stripped = text.strip()
    return (
        text_stripped.endswith(':') and
        any(keyword in text_stripped for keyword in ['Cost Base', 'Capital Gain', 'Capital Proceeds', 'Requirements'])
    )


def add_calculation_line(doc, text, is_step=False, is_result=False, is_header=False):
    """Add a single calculation line with appropriate formatting"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(0)

    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(11)

    if is_result:
        run.bold = True
        run.font.color.rgb = RGBColor(0, 100, 0)  # Dark green
    elif is_step:
        run.bold = True
        run.font.color.rgb = RGBColor(0, 51, 102)  # Dark blue
    elif is_header:
        run.bold = True


def convert_markdown_to_docx(md_path, docx_path):
    """Main conversion function"""
    print(f"Reading markdown file: {md_path}")

    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()

    lines = content.split('\n')
    doc = create_styled_document()

    i = 0
    in_calculation = False
    in_toc = False  # Table of contents section

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Handle empty lines
        if not stripped:
            if in_calculation:
                # Empty line might end calculation block
                # Check if next non-empty line is a heading or separator
                next_idx = i + 1
                while next_idx < len(lines) and not lines[next_idx].strip():
                    next_idx += 1
                if next_idx < len(lines):
                    next_line = lines[next_idx].strip()
                    if next_line.startswith('#') or next_line == '---' or next_line.startswith('|'):
                        in_calculation = False
            i += 1
            continue

        # Handle horizontal rules
        if stripped == '---':
            in_calculation = False
            add_horizontal_line(doc)
            i += 1
            continue

        # Handle H1 headings
        if stripped.startswith('# ') and not stripped.startswith('## '):
            in_calculation = False
            in_toc = False
            title = stripped[2:]
            p = doc.add_heading(title, level=0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue

        # Handle H2 headings (Scenario titles)
        if stripped.startswith('## '):
            in_calculation = False
            title = stripped[3:]
            if 'Table of Contents' in title:
                in_toc = True
            else:
                in_toc = False
            doc.add_heading(title, level=1)
            i += 1
            continue

        # Handle H3 headings (Subsections)
        if stripped.startswith('### '):
            in_calculation = False
            in_toc = False
            title = stripped[4:]

            # Check if this is a CGT Calculation header
            if 'CGT Calculation' in title:
                doc.add_heading(title, level=2)
                in_calculation = True
            else:
                doc.add_heading(title, level=2)
            i += 1
            continue

        # Handle tables
        if stripped.startswith('|'):
            in_calculation = False
            in_toc = False
            table_data, new_idx = parse_table(lines, i)
            if table_data:
                add_table_to_doc(doc, table_data)
            i = new_idx
            continue

        # Handle Table of Contents numbered items with links
        if in_toc and re.match(r'^\d+\.\s*\[', stripped):
            # Extract text from markdown link
            match = re.match(r'^(\d+)\.\s*\[([^\]]+)\]', stripped)
            if match:
                num = match.group(1)
                text = match.group(2)
                p = doc.add_paragraph()
                run = p.add_run(f"{num}. {text}")
                run.font.name = 'Calibri'
                run.font.size = Pt(11)
                i += 1
                continue

        # Handle calculation blocks
        if in_calculation or is_step_line(stripped):
            in_calculation = True

            # Check for Step lines
            if is_step_line(stripped):
                add_calculation_line(doc, stripped, is_step=True)
                i += 1
                continue

            # Check for RESULT lines
            if is_result_line(stripped):
                add_calculation_line(doc, stripped, is_result=True)
                i += 1
                continue

            # Check for calculation headers (Cost Base:, etc.)
            if is_calculation_header(stripped):
                add_calculation_line(doc, stripped, is_header=True)
                i += 1
                continue

            # Regular calculation line
            add_calculation_line(doc, stripped)
            i += 1
            continue

        # Handle bold property/section headers (like **Property 1 - Beach House**)
        if stripped.startswith('**') and stripped.endswith('**'):
            in_toc = False
            text = stripped[2:-2]
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            run = p.add_run(text)
            run.bold = True
            run.font.size = Pt(12)
            run.font.name = 'Calibri'
            run.font.color.rgb = RGBColor(0, 51, 102)
            i += 1
            continue

        # Handle bullet points
        if stripped.startswith('- ') or stripped.startswith('* '):
            in_toc = False
            text = stripped[2:]
            p = doc.add_paragraph(style='List Bullet')
            process_inline_formatting(p, text)
            i += 1
            continue

        # Handle numbered lists (cost base elements, etc.)
        if re.match(r'^\d+\.\s', stripped):
            in_toc = False
            text = re.sub(r'^\d+\.\s', '', stripped)
            num_match = re.match(r'^(\d+)\.\s', stripped)
            num = num_match.group(1) if num_match else '•'
            p = doc.add_paragraph()
            run_num = p.add_run(f"{num}. ")
            run_num.font.name = 'Calibri'
            run_num.font.size = Pt(11)
            run_num.bold = True
            process_inline_formatting(p, text)
            i += 1
            continue

        # Handle URLs (standalone URLs)
        if stripped.startswith('http://') or stripped.startswith('https://'):
            p = doc.add_paragraph()
            run = p.add_run(stripped)
            run.font.name = 'Calibri'
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0, 102, 204)
            i += 1
            continue

        # Regular paragraph
        p = doc.add_paragraph()
        process_inline_formatting(p, stripped)
        i += 1

    # Save document
    print(f"Saving Word document: {docx_path}")
    doc.save(docx_path)
    print("Conversion complete!")

    # Print summary
    print(f"\nDocument saved successfully to:\n{docx_path}")


if __name__ == '__main__':
    import os

    # Get the directory of this script
    script_dir = os.path.dirname(os.path.abspath(__file__))

    md_file = os.path.join(script_dir, 'ALL_CGT_SCENARIOS.md')
    docx_file = os.path.join(script_dir, 'ALL_CGT_SCENARIOS.docx')

    convert_markdown_to_docx(md_file, docx_file)
